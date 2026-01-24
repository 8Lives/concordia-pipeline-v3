"""
Concordia Pipeline v3 - Streamlit Application

RAG-enhanced clinical trial data harmonization with:
- Dynamic specification retrieval
- Real-time progress tracking
- Interactive QC review
- Full traceability

Run with: streamlit run app.py

Secrets Configuration:
    Create .streamlit/secrets.toml with:
        HUGGINGFACE_TOKEN = "hf_xxx..."
        ANTHROPIC_API_KEY = "sk-ant-xxx..."  # Optional, for LLM features
        VOYAGE_API_KEY = "pa-xxx..."  # Optional, for Voyage embeddings
"""

import streamlit as st
import pandas as pd
import os
import time
from pathlib import Path
from datetime import datetime
from typing import Dict
import json
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def load_secrets_to_env():
    """Load Streamlit secrets into environment variables."""
    secrets_loaded = []

    # Map of secret names to environment variable names
    secret_mappings = {
        "HUGGINGFACE_TOKEN": ["HUGGINGFACE_TOKEN", "HF_TOKEN"],
        "HF_TOKEN": ["HUGGINGFACE_TOKEN", "HF_TOKEN"],
        "ANTHROPIC_API_KEY": ["ANTHROPIC_API_KEY"],
        "VOYAGE_API_KEY": ["VOYAGE_API_KEY"],
    }

    try:
        for secret_name, env_names in secret_mappings.items():
            if secret_name in st.secrets:
                value = st.secrets[secret_name]
                for env_name in env_names:
                    os.environ[env_name] = value
                secrets_loaded.append(secret_name)
                logger.info(f"Loaded secret: {secret_name}")
    except Exception as e:
        logger.warning(f"Could not load secrets: {e}")

    return secrets_loaded


# Load secrets before importing pipeline components
secrets_loaded = load_secrets_to_env()

# Page configuration
st.set_page_config(
    page_title="Concordia Pipeline v3",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Import pipeline components
try:
    from orchestrator import PipelineOrchestrator, create_orchestrator, PipelineResult
    from config.settings import get_settings, reset_settings
    PIPELINE_AVAILABLE = True
except ImportError as e:
    PIPELINE_AVAILABLE = False
    IMPORT_ERROR = str(e)


def parse_data_dictionary(df: pd.DataFrame) -> Dict:
    """
    Parse a data dictionary DataFrame using deterministic rules (v2 approach).

    This uses the proven parsing logic from v2 that correctly handles
    clinical data dictionaries with continuation rows for code values.

    Output format:
    {
        "SEX": {"codes": {"1": "Male", "2": "Female"}},
        "RACE": {"codes": {"11": "White", "12": "Black or African American"}}
    }
    """
    dictionary = {}

    # Find the header row by looking for known column headers
    header_row_idx = None
    for idx, row in df.iterrows():
        for col_idx, cell in enumerate(row):
            if pd.notna(cell):
                cell_str = str(cell).strip().upper()
                if cell_str in ['VARIABLE NAME', 'VARIABLE', 'VAR NAME']:
                    header_row_idx = idx
                    break
        if header_row_idx is not None:
            break

    # If we found a header row, reindex the DataFrame
    if header_row_idx is not None:
        new_columns = df.iloc[header_row_idx].tolist()
        df = df.iloc[header_row_idx + 1:].copy()
        df.columns = [str(c).strip() if pd.notna(c) else f'col_{i}' for i, c in enumerate(new_columns)]

    # Identify key columns by mapping column names (case-insensitive)
    col_map = {str(c).upper().replace('\n', ' '): c for c in df.columns}

    var_col = None
    value_col = None
    format_col = None

    # Find variable name column
    for candidate in ['VARIABLE NAME', 'VARIABLE', 'VAR', 'NAME', 'FIELD']:
        if candidate in col_map:
            var_col = col_map[candidate]
            break

    # Find value/decode column
    for candidate in ['VALID VALUES', 'VALUES', 'DECODE', 'VALID VALUE']:
        if candidate in col_map:
            value_col = col_map[candidate]
            break

    # Find format column
    for candidate in ['FORMAT  (VALUE LIST)', 'FORMAT (VALUE LIST)', 'FORMAT', 'VALUE LIST', 'CODELIST']:
        if candidate in col_map:
            format_col = col_map[candidate]
            break

    if not var_col:
        logger.warning("Could not find variable name column in dictionary")
        return {}

    # Parse the dictionary - track current variable for continuation rows
    current_var = None

    for _, row in df.iterrows():
        var_name = row.get(var_col)

        # If we have a variable name, start tracking a new variable
        if pd.notna(var_name) and str(var_name).strip():
            current_var = str(var_name).strip().upper()
            if current_var not in dictionary:
                dictionary[current_var] = {
                    "codes": {},
                    "format": str(row.get(format_col, '')) if format_col and pd.notna(row.get(format_col)) else ''
                }

        # Parse code = value patterns from the value column
        if current_var and value_col:
            value_str = row.get(value_col)
            if pd.notna(value_str):
                value_str = str(value_str).strip()
                if '=' in value_str:
                    parts = value_str.split('=', 1)
                    code = parts[0].strip()
                    label = parts[1].strip()
                    dictionary[current_var]["codes"][code] = label

    # Remove variables with no codes
    dictionary = {k: v for k, v in dictionary.items() if v.get("codes")}

    logger.info(f"Deterministic parser found {len(dictionary)} variables with codes")
    for var, data in dictionary.items():
        codes = data.get("codes", {})
        logger.info(f"  {var}: {len(codes)} codes - keys: {list(codes.keys())}")

    return dictionary


def init_session_state():
    """Initialize session state variables."""
    if "pipeline_result" not in st.session_state:
        st.session_state.pipeline_result = None
    if "progress_log" not in st.session_state:
        st.session_state.progress_log = []
    if "current_stage" not in st.session_state:
        st.session_state.current_stage = None
    if "orchestrator" not in st.session_state:
        st.session_state.orchestrator = None


def progress_callback(stage: str, status: str, message: str, progress: float):
    """Legacy callback for pipeline progress updates (kept for compatibility)."""
    st.session_state.current_stage = stage
    st.session_state.progress_log.append({
        "timestamp": datetime.now().isoformat(),
        "stage": stage,
        "status": status,
        "message": message,
        "progress": progress
    })


def render_sidebar():
    """Render the sidebar with configuration options."""
    with st.sidebar:
        st.title("⚙️ Configuration")

        st.subheader("RAG Settings")
        use_rag = st.checkbox(
            "Enable RAG",
            value=True,
            help="Use RAG for dynamic specification retrieval"
        )

        embedding_provider = st.selectbox(
            "Embedding Provider",
            ["local", "mock"],
            index=0,
            help="Local uses sentence-transformers, mock for testing"
        )

        st.subheader("LLM Settings")
        use_llm = st.checkbox(
            "Enable LLM",
            value=True,
            help="Use Claude for value resolution and review"
        )

        enable_review = st.checkbox(
            "Enable LLM Review",
            value=True,
            disabled=not use_llm,
            help="Run LLM-powered quality review stage"
        )

        st.subheader("Pipeline Options")
        skip_qc = st.checkbox(
            "Skip QC Stage",
            value=False,
            help="Skip quality control checks"
        )

        st.divider()

        st.subheader("About")
        st.markdown("""
        **Concordia Pipeline v3**

        RAG-enhanced harmonization for
        SDTM Demographics (DM) domain.

        Features:
        - 🔍 Dynamic spec retrieval
        - 🤖 LLM-powered resolution
        - 📊 18 output variables
        - ✅ Automated QC checks
        - 🔎 LLM Review stage
        - 📝 Full lineage tracking
        """)

        # Show secrets status
        st.divider()
        st.subheader("🔑 API Keys")

        hf_configured = bool(os.environ.get("HUGGINGFACE_TOKEN") or os.environ.get("HF_TOKEN"))
        anthropic_configured = bool(os.environ.get("ANTHROPIC_API_KEY"))
        voyage_configured = bool(os.environ.get("VOYAGE_API_KEY"))

        st.markdown(f"- HuggingFace: {'✅' if hf_configured else '❌'}")
        st.markdown(f"- Anthropic: {'✅' if anthropic_configured else '⚪ (optional)'}")
        st.markdown(f"- Voyage AI: {'✅' if voyage_configured else '⚪ (optional)'}")

        if not hf_configured:
            st.caption("Add HUGGINGFACE_TOKEN to .streamlit/secrets.toml")

        return {
            "use_rag": use_rag,
            "use_llm": use_llm,
            "enable_review": enable_review,
            "embedding_provider": embedding_provider,
            "skip_qc": skip_qc
        }


def render_file_upload():
    """Render file upload section."""
    st.subheader("📁 Input Data")

    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded_file = st.file_uploader(
            "Upload source data file",
            type=["csv", "xlsx", "xls", "sas7bdat"],
            help="Supported formats: CSV, Excel, SAS"
        )

    with col2:
        trial_id = st.text_input(
            "Trial ID (optional)",
            placeholder="e.g., NCT12345678",
            help="Leave blank to extract from filename"
        )

    # Optional data dictionary upload
    with st.expander("📖 Data Dictionary (Optional)", expanded=False):
        st.caption("Upload a data dictionary to improve column mapping accuracy")
        data_dict_file = st.file_uploader(
            "Upload data dictionary",
            type=["csv", "xlsx", "xls", "json"],
            help="Maps source column names to descriptions or standard terms",
            key="data_dict_uploader"
        )

        if data_dict_file:
            st.success(f"✅ Data dictionary loaded: {data_dict_file.name}")

    return uploaded_file, trial_id, data_dict_file if 'data_dict_file' in dir() else None


def render_progress():
    """Render progress indicators."""
    if st.session_state.progress_log:
        latest = st.session_state.progress_log[-1]

        # Progress bar
        progress_bar = st.progress(latest["progress"])

        # Status message
        stage_emoji = {
            "init": "🔧",
            "ingest": "📥",
            "map": "🗺️",
            "harmonize": "✨",
            "qc": "✅",
            "review": "🔎",
            "finalize": "📦",
            "complete": "🎉",
            "error": "❌"
        }
        emoji = stage_emoji.get(latest["stage"], "⏳")
        st.info(f"{emoji} **{latest['stage'].upper()}**: {latest['message']}")


def create_results_zip(result: PipelineResult) -> bytes:
    """Create a ZIP file containing all pipeline results.

    Includes:
    - Harmonized data CSV
    - QC report CSV
    - Transformation report DOCX (with row-by-row details)

    Excludes JSON files per user request.
    """
    import io
    import zipfile

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        trial_id = result.metadata.get("trial_id", "output")

        # Add harmonized data
        if result.harmonized_data is not None:
            csv_data = result.harmonized_data.to_csv(index=False)
            zf.writestr(f"{trial_id}_harmonized.csv", csv_data)

        # Add QC report
        if result.qc_report is not None and len(result.qc_report) > 0:
            qc_csv = result.qc_report.to_csv(index=False)
            zf.writestr(f"{trial_id}_qc_report.csv", qc_csv)

        # Add transformation report as DOCX (v2 format with row-by-row summary)
        try:
            docx_report = create_transformation_report_docx(result)
            zf.writestr(f"{trial_id}_transformation_report.docx", docx_report)
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX report")
        except Exception as e:
            logger.warning(f"Could not create DOCX report: {e}")

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def create_transformation_report_docx(result: PipelineResult) -> bytes:
    """
    Create a DOCX transformation report matching v2 format exactly.

    Returns bytes for the DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import io

    doc = Document()

    # Title (v2 format)
    title = doc.add_heading('Harmonization Transformation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get metadata
    ingest_meta = result.metadata.get('stages', {}).get('ingest', {})
    harmonize_meta = result.metadata.get('stages', {}).get('harmonize', {})

    trial_id = result.metadata.get('trial_id', 'Unknown')
    source_filename = ingest_meta.get('source_filename', result.metadata.get('source_filename', 'Unknown'))
    rows_in = ingest_meta.get('rows', result.metadata.get('rows_processed', 0))
    rows_out = harmonize_meta.get('rows_out', rows_in)
    rows_dropped = rows_in - rows_out if rows_in and rows_out else 0

    # Header info (v2 format - not a table, just paragraphs)
    doc.add_paragraph(f"Trial: {trial_id}")
    doc.add_paragraph(f"Input: {source_filename}")
    doc.add_paragraph(f"Rows in input: {rows_in} | Rows in output: {rows_out} | Rows dropped: {rows_dropped}")
    doc.add_paragraph(f"Pipeline version: v3 (RAG-Enhanced Agentic Architecture)")
    doc.add_paragraph(f"Run ID: {result.metadata.get('timestamp', 'Unknown')}")

    # Execution Summary (v2 format)
    doc.add_heading('Execution Summary', level=1)
    stages = result.metadata.get('stages', {})
    for stage_name, stage_info in stages.items():
        if isinstance(stage_info, dict):
            status = "✓" if stage_info.get('success', True) else "✗"
            time_ms = stage_info.get('execution_time_ms', 0)
            doc.add_paragraph(f"{status} {stage_name}: {time_ms}ms")

    # Dictionary info
    dict_filename = ingest_meta.get('dictionary_filename') or result.metadata.get('dictionary_filename')
    if dict_filename:
        doc.add_paragraph(f"Dictionary used: {dict_filename}")
    else:
        doc.add_paragraph("Dictionary used: None")

    doc.add_paragraph()

    # Output Schema (v2 format)
    doc.add_heading('1. Output Schema', level=1)
    if result.harmonized_data is not None:
        doc.add_paragraph(', '.join(result.harmonized_data.columns.tolist()))
    else:
        doc.add_paragraph("TRIAL, SUBJID, SEX, RACE, AGE, AGEU, AGEGP, ETHNIC, COUNTRY, SITEID, STUDYID, USUBJID, ARMCD, ARM, BRTHDTC, RFSTDTC, RFENDTC, DOMAIN")

    # Variable Transformation Table (v2 format)
    doc.add_heading('2. Variable-Level Transformations', level=1)

    if result.lineage:
        trans_table = doc.add_table(rows=1, cols=7)
        trans_table.style = 'Table Grid'

        # Headers (v2 format)
        headers = ['Variable', 'Source', 'Operation', 'Details', 'Changed', '%', 'Missing']
        hdr_cells = trans_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Data rows
        for entry in result.lineage:
            row_cells = trans_table.add_row().cells
            row_cells[0].text = str(entry.get('variable', ''))

            # Source column - show (derived) if None
            source_col = entry.get('source_column', '') or ''
            row_cells[1].text = str(source_col) if source_col else '(derived)'

            row_cells[2].text = str(entry.get('mapping_operation', 'Copy'))

            # Use transform_operation (v2 key name), truncated to 50 chars
            transform_op = entry.get('transform_operation', '') or entry.get('transformation', 'None')
            row_cells[3].text = str(transform_op)[:50]

            row_cells[4].text = str(entry.get('rows_changed', 0))
            row_cells[5].text = f"{entry.get('percent_changed', 0):.1f}%"
            row_cells[6].text = str(entry.get('missing_count', 0))

    # QC Report Section (v2 format)
    doc.add_heading('3. QC Report', level=1)

    if result.qc_report is not None and len(result.qc_report) > 0:
        qc_table = doc.add_table(rows=1, cols=5)
        qc_table.style = 'Table Grid'

        qc_headers = ['TRIAL', 'Issue Type', 'Variable', 'Rows Affected', 'Notes']
        hdr_cells = qc_table.rows[0].cells
        for i, header in enumerate(qc_headers):
            hdr_cells[i].text = header

        for _, issue in result.qc_report.iterrows():
            row_cells = qc_table.add_row().cells
            row_cells[0].text = str(trial_id)
            row_cells[1].text = str(issue.get('issue_type', ''))
            row_cells[2].text = str(issue.get('variable', ''))
            row_cells[3].text = str(issue.get('n_rows_affected', 0))
            row_cells[4].text = str(issue.get('notes', ''))[:100]
    else:
        doc.add_paragraph("No QC issues found.")

    # Files Produced (v2 format)
    doc.add_heading('4. Files Produced', level=1)
    doc.add_paragraph(f"Harmonized output: {trial_id}_DM_harmonized_*.csv")
    doc.add_paragraph(f"QC report: {trial_id}_QC_report_*.csv")
    doc.add_paragraph(f"Transformation report: This document")

    # LLM Review section (v3 addition)
    if result.review_result:
        doc.add_heading('5. LLM Review', level=1)

        review = result.review_result
        doc.add_paragraph(f"Approval: {review.get('approval', 'Unknown')}")
        doc.add_paragraph(f"Quality: {review.get('overall_quality', 'Unknown')}")

        if review.get('reason'):
            doc.add_paragraph(f"Summary: {review.get('reason')}")

    # Warnings and Errors
    if result.warnings:
        doc.add_heading('Warnings', level=1)
        for w in result.warnings:
            doc.add_paragraph(f"• {w}", style='List Bullet')

    if result.errors:
        doc.add_heading('Errors', level=1)
        for e in result.errors:
            doc.add_paragraph(f"• {e}", style='List Bullet')

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_results(result: PipelineResult):
    """Render pipeline results."""
    st.header("📊 Results")

    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric(
            "Status",
            "✅ Success" if result.success else "❌ Failed"
        )

    with col2:
        rows = len(result.harmonized_data) if result.harmonized_data is not None else 0
        st.metric("Rows Processed", rows)

    with col3:
        qc_issues = len(result.qc_report) if result.qc_report is not None else 0
        st.metric("QC Issues", qc_issues)

    with col4:
        time_sec = result.execution_time_ms / 1000
        st.metric("Time", f"{time_sec:.1f}s")

    # LLM metrics row
    col5, col6, col7, col8 = st.columns(4)

    with col5:
        llm_enabled = result.metadata.get("llm_enabled", False)
        st.metric("LLM", "✅ Enabled" if llm_enabled else "⚪ Disabled")

    with col6:
        tokens = result.metadata.get("llm_tokens_used", 0)
        st.metric("LLM Tokens", f"{tokens:,}" if tokens else "-")

    with col7:
        review_status = "N/A"
        if result.review_result:
            approval = result.review_result.get("approval", "unknown")
            review_status = approval.replace("_", " ").title()
        st.metric("Review", review_status)

    with col8:
        model = result.metadata.get("llm_model", "none")
        st.metric("Model", model if model != "none" else "-")

    # Errors and warnings
    if result.errors:
        st.error("**Errors:**\n" + "\n".join(f"- {e}" for e in result.errors))

    if result.warnings:
        st.warning("**Warnings:**\n" + "\n".join(f"- {w}" for w in result.warnings))

    # Download All as ZIP button
    if result.success:
        zip_data = create_results_zip(result)
        if zip_data:
            trial_id = result.metadata.get("trial_id", "output")
            st.download_button(
                "📦 Download All Results (ZIP)",
                zip_data,
                file_name=f"{trial_id}_harmonization_results.zip",
                mime="application/zip",
                type="primary"
            )

    # Tabs for detailed results (matching v2 structure)
    tabs = st.tabs(["📋 Harmonized Data", "⚠️ QC Report", "🔎 LLM Review", "🔄 Transformations", "📦 Downloads"])

    with tabs[0]:
        if result.harmonized_data is not None:
            st.dataframe(result.harmonized_data, use_container_width=True)

            # Download button
            csv = result.harmonized_data.to_csv(index=False)
            st.download_button(
                "⬇️ Download Harmonized Data",
                csv,
                file_name="harmonized_data.csv",
                mime="text/csv"
            )
        else:
            st.info("No harmonized data available")

    with tabs[1]:
        if result.qc_report is not None and len(result.qc_report) > 0:
            st.dataframe(result.qc_report, use_container_width=True)

            csv = result.qc_report.to_csv(index=False)
            st.download_button(
                "⬇️ Download QC Report",
                csv,
                file_name="qc_report.csv",
                mime="text/csv"
            )
        else:
            st.success("No QC issues found! ✨")

    with tabs[2]:
        if result.review_result:
            review = result.review_result

            # Show STOPLIGHT status (new v3 format)
            stoplight = review.get("stoplight") or review.get("approval", "unknown")

            if stoplight.upper() == "GREEN":
                st.success(f"🟢 **STOPLIGHT: GREEN**")
                st.markdown("All 5 core variables present and properly formatted.")
            elif stoplight.upper() == "YELLOW":
                st.warning(f"🟡 **STOPLIGHT: YELLOW**")
                st.markdown("Missing 1-2 core variables or formatting issues detected.")
            elif stoplight.upper() == "RED":
                st.error(f"🔴 **STOPLIGHT: RED**")
                st.markdown("Missing 3+ core variables.")
            else:
                st.info(f"🔍 **LLM Review: {stoplight.upper()}**")

            # Show core variables status
            if "core_variables_present" in review:
                st.markdown(f"**Core Variables Present:** {', '.join(review['core_variables_present'])}")
            if "core_variables_missing" in review:
                missing = review['core_variables_missing']
                if missing:
                    st.markdown(f"**Core Variables Missing:** {', '.join(missing)}")

            # Show quality assessment
            if "overall_quality" in review:
                st.markdown(f"**Quality Rating:** {review['overall_quality']}")

            # Show reason
            if "reason" in review:
                st.markdown(f"**Summary:** {review['reason']}")

            # Show formatting issues
            if review.get("formatting_issues"):
                st.subheader("Formatting Issues")
                for issue in review["formatting_issues"]:
                    st.markdown(f"- {issue}")

            # Show critical issues
            if review.get("critical_issues"):
                st.subheader("Critical Issues")
                for issue in review["critical_issues"]:
                    if isinstance(issue, dict):
                        st.markdown(f"- {issue.get('issue', issue)}")
                    else:
                        st.markdown(f"- {issue}")

            # Show recommendations
            if review.get("recommendations"):
                st.subheader("Recommendations")
                for rec in review["recommendations"]:
                    st.markdown(f"- {rec}")

            # Full details
            with st.expander("Full Review Details"):
                st.json(review)
        else:
            st.info("No LLM review performed (requires Anthropic API key)")

    with tabs[3]:
        st.subheader("Variable Transformations")
        if result.lineage:
            # Convert lineage to DataFrame with v2 column names
            lineage_data = []
            for entry in result.lineage:
                lineage_data.append({
                    "variable": entry.get("variable", ""),
                    "source_column": entry.get("source_column", "") or "(derived)",
                    "mapping_operation": entry.get("mapping_operation", ""),
                    "transform_operation": entry.get("transform_operation", "") or entry.get("transformation", ""),
                    "transform_details": str(entry.get("transform_details", {})),
                    "rows_changed": entry.get("rows_changed", 0),
                    "percent_changed": f"{entry.get('percent_changed', 0):.1f}%",
                    "missing_count": entry.get("missing_count", 0),
                    "non_null_count": entry.get("non_null_count", 0)
                })

            lineage_df = pd.DataFrame(lineage_data)
            st.dataframe(lineage_df, use_container_width=True)
        else:
            st.info("No transformation data available")

    with tabs[4]:
        st.subheader("Downloads")

        # Individual downloads
        col1, col2 = st.columns(2)

        with col1:
            if result.harmonized_data is not None:
                trial_id = result.metadata.get("trial_id", "output")
                csv = result.harmonized_data.to_csv(index=False)
                st.download_button(
                    "📋 Harmonized Data (CSV)",
                    csv,
                    file_name=f"{trial_id}_harmonized.csv",
                    mime="text/csv"
                )

            if result.qc_report is not None and len(result.qc_report) > 0:
                qc_csv = result.qc_report.to_csv(index=False)
                st.download_button(
                    "⚠️ QC Report (CSV)",
                    qc_csv,
                    file_name=f"{trial_id}_qc_report.csv",
                    mime="text/csv"
                )

        with col2:
            # DOCX Transformation Report
            try:
                docx_bytes = create_transformation_report_docx(result)
                st.download_button(
                    "📄 Transformation Report (DOCX)",
                    docx_bytes,
                    file_name=f"{trial_id}_transformation_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.warning(f"Could not generate DOCX report: {e}")

            # All as ZIP
            if result.success:
                zip_data = create_results_zip(result)
                if zip_data:
                    st.download_button(
                        "📦 All Results (ZIP)",
                        zip_data,
                        file_name=f"{trial_id}_harmonization_results.zip",
                        mime="application/zip"
                    )

    # Metadata
    with st.expander("📋 Metadata"):
        st.json(result.metadata)


def run_pipeline(uploaded_file, trial_id: str, config: dict, data_dict_file=None):
    """Run the harmonization pipeline."""
    # Reset progress
    st.session_state.progress_log = []
    st.session_state.pipeline_result = None

    # Read data dictionary if provided - we'll parse it with LLM after creating orchestrator
    dict_df = None
    data_dict = None
    if data_dict_file is not None:
        try:
            if data_dict_file.name.endswith('.csv'):
                dict_df = pd.read_csv(data_dict_file)
            elif data_dict_file.name.endswith(('.xlsx', '.xls')):
                dict_df = pd.read_excel(data_dict_file)
            elif data_dict_file.name.endswith('.json'):
                # JSON is already structured, use directly
                data_dict = json.load(data_dict_file)
            logger.info(f"Loaded dictionary file: {data_dict_file.name}")
        except Exception as e:
            st.warning(f"Could not load data dictionary: {e}")
            logger.exception("Data dictionary load error")

    # Read uploaded file
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
        elif uploaded_file.name.endswith('.sas7bdat'):
            # SAS7BDAT format requires pyreadstat or sas7bdat package
            try:
                import pyreadstat
                df, meta = pyreadstat.read_sas7bdat(uploaded_file)
                logger.info(f"Loaded SAS file with {len(df)} rows using pyreadstat")
            except ImportError:
                try:
                    from sas7bdat import SAS7BDAT
                    with SAS7BDAT(uploaded_file) as f:
                        df = f.to_data_frame()
                    logger.info(f"Loaded SAS file with {len(df)} rows using sas7bdat")
                except ImportError:
                    st.error("SAS7BDAT support requires 'pyreadstat' or 'sas7bdat' package. "
                             "Install with: pip install pyreadstat")
                    return None
        else:
            st.error(f"Unsupported file format: {uploaded_file.name}")
            return None
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

    # Extract trial_id from filename if not provided
    if not trial_id:
        import re
        match = re.search(r'(NCT\d{8})', uploaded_file.name)
        trial_id = match.group(1) if match else Path(uploaded_file.name).stem

    # Create progress display elements
    progress_container = st.container()
    with progress_container:
        progress_bar = st.progress(0, text="Initializing pipeline...")
        status_text = st.empty()

    # Create a progress callback that updates Streamlit elements
    def streamlit_progress_callback(stage: str, status: str, message: str, progress: float):
        """Update Streamlit progress elements in real-time."""
        stage_emoji = {
            "init": "🔧",
            "ingest": "📥",
            "map": "🗺️",
            "harmonize": "✨",
            "qc": "✅",
            "review": "🔎",
            "finalize": "📦",
            "complete": "🎉",
            "error": "❌"
        }
        emoji = stage_emoji.get(stage, "⏳")

        # Update progress bar (clamp to 0-1 range)
        progress_value = max(0.0, min(1.0, progress))
        progress_bar.progress(progress_value, text=f"{emoji} {stage.upper()}: {message}")

        # Also log to session state for history
        st.session_state.progress_log.append({
            "timestamp": datetime.now().isoformat(),
            "stage": stage,
            "status": status,
            "message": message,
            "progress": progress
        })

    # Create orchestrator with real-time progress callback
    reset_settings()
    orchestrator = create_orchestrator(
        use_rag=config["use_rag"],
        use_llm=config["use_llm"],
        enable_review=config["enable_review"],
        progress_callback=streamlit_progress_callback,
        embedding_provider=config["embedding_provider"]
    )

    # Parse data dictionary using deterministic parser (v2 approach)
    if dict_df is not None and data_dict is None:
        progress_bar.progress(0.05, text="🔧 INIT: Parsing data dictionary...")

        try:
            data_dict = parse_data_dictionary(dict_df)
            if data_dict:
                st.success(f"✅ Dictionary parsed: {len(data_dict)} variables extracted")
                logger.info(f"Parsed dictionary variables: {list(data_dict.keys())}")
                for var, info in data_dict.items():
                    logger.info(f"  {var}: {list(info.get('codes', {}).keys())}")
            else:
                st.warning("Could not extract code mappings from dictionary")
        except Exception as e:
            logger.exception("Failed to parse dictionary")
            st.warning(f"Dictionary parsing failed: {e}")

    # Run pipeline (no spinner - we have the progress bar)
    try:
        result = orchestrator.run(
            input_df=df,
            trial_id=trial_id,
            skip_qc=config["skip_qc"],
            data_dict=data_dict
        )
    except Exception as e:
        logger.exception("Pipeline run failed with exception")
        st.error(f"Pipeline error: {str(e)}")
        import traceback
        with st.expander("Error Details"):
            st.code(traceback.format_exc())
        return None

    # Clear progress elements after completion
    if result.success:
        progress_bar.progress(1.0, text="🎉 Pipeline complete!")
    else:
        progress_bar.progress(1.0, text="❌ Pipeline failed")
        if result.errors:
            for err in result.errors:
                st.error(f"Error: {err}")

    st.session_state.pipeline_result = result
    return result


def main():
    """Main application entry point."""
    init_session_state()

    # Header
    st.title("🔬 Concordia Pipeline v3")
    st.markdown("*RAG-Enhanced Clinical Trial Data Harmonization*")

    if not PIPELINE_AVAILABLE:
        st.error(f"Pipeline components not available: {IMPORT_ERROR}")
        st.info("Please ensure all dependencies are installed.")
        return

    # Sidebar configuration
    config = render_sidebar()

    # Main content area
    col1, col2 = st.columns([1, 1])

    with col1:
        uploaded_file, trial_id, data_dict_file = render_file_upload()

        if uploaded_file is not None:
            # Show file preview
            with st.expander("📄 File Preview", expanded=True):
                try:
                    if uploaded_file.name.endswith('.csv'):
                        preview_df = pd.read_csv(uploaded_file, nrows=5)
                    elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                        preview_df = pd.read_excel(uploaded_file, nrows=5)
                    elif uploaded_file.name.endswith('.sas7bdat'):
                        # For SAS files, read full file then take head (pyreadstat doesn't support nrows)
                        try:
                            import pyreadstat
                            preview_df, _ = pyreadstat.read_sas7bdat(uploaded_file)
                            preview_df = preview_df.head(5)
                        except ImportError:
                            try:
                                from sas7bdat import SAS7BDAT
                                with SAS7BDAT(uploaded_file) as f:
                                    preview_df = f.to_data_frame().head(5)
                            except ImportError:
                                st.warning("Install 'pyreadstat' to preview SAS files")
                                preview_df = None
                    else:
                        preview_df = None

                    if preview_df is not None:
                        st.dataframe(preview_df, use_container_width=True)
                        st.caption(f"Showing first 5 rows • {len(preview_df.columns)} columns")

                    # Reset file position for later reading
                    uploaded_file.seek(0)
                except Exception as e:
                    st.error(f"Error previewing file: {e}")

    with col2:
        st.subheader("🚀 Run Pipeline")

        if uploaded_file is not None:
            if st.button("▶️ Start Harmonization", type="primary", use_container_width=True):
                result = run_pipeline(uploaded_file, trial_id, config, data_dict_file)
        else:
            st.info("Upload a file to begin")

    # Results section
    st.divider()

    if st.session_state.pipeline_result is not None:
        render_results(st.session_state.pipeline_result)


if __name__ == "__main__":
    main()
